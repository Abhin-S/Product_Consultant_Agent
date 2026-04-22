from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
import re

import docx
import pdfplumber

from config import settings


logger = logging.getLogger(__name__)


@dataclass
class Document:
    text: str
    source: str
    metadata: dict[str, str | int | float | bool] = field(default_factory=dict)


def _table_to_markdown(table: list[list[str | None]]) -> str:
    rows = [[(cell or "").replace("\n", " ").strip() for cell in row] for row in table if row]
    if len(rows) < 2:
        return ""

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]

    header = normalized[0]
    if all(not cell for cell in header) and len(normalized) > 1:
        header = normalized[1]
        normalized = normalized[1:]

    body = normalized[1:] if len(normalized) > 1 else []
    if all(not cell for cell in header):
        return ""

    lines = [
        "| " + " | ".join(cell.replace("|", "\\|") for cell in header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |" for row in body)
    return "\n".join(lines)


def _safe_extract_page_text(page: pdfplumber.page.Page, *, context: str) -> str:
    try:
        return page.extract_text() or ""
    except Exception as exc:
        logger.warning("PDF text extraction failed (%s): %s", context, exc)
        return ""


def _extract_text_outside_tables(
    page: pdfplumber.page.Page,
    table_bboxes: list[tuple],
    *,
    context: str,
) -> str:
    if not table_bboxes:
        return _safe_extract_page_text(page, context=context)

    try:
        def _outside(obj: dict) -> bool:
            ox = obj.get("x0", 0)
            oy = obj.get("top", 0)
            for tx0, ty0, tx1, ty1 in table_bboxes:
                if tx0 - 2 <= ox <= tx1 + 2 and ty0 - 2 <= oy <= ty1 + 2:
                    return False
            return True

        filtered = page.filter(_outside)
        try:
            return filtered.extract_text() or ""
        except Exception as exc:
            logger.warning("Filtered PDF text extraction failed (%s): %s", context, exc)
            return _safe_extract_page_text(page, context=context)
    except Exception as exc:
        logger.warning("Table-region filtering failed (%s): %s", context, exc)
        return _safe_extract_page_text(page, context=context)


def _derive_source_metadata(path: Path, root: Path) -> dict[str, str | int | float | bool]:
    rel_path = path.resolve().relative_to(root.resolve())
    parts = list(rel_path.parts)

    topic = parts[0] if len(parts) > 1 else "General"
    subtopic = parts[1] if len(parts) > 2 else ""

    return {
        "source_rel": rel_path.as_posix(),
        "source_name": path.name,
        "source_stem": path.stem,
        "source_ext": path.suffix.lower(),
        "topic": topic,
        "subtopic": subtopic,
    }


def _load_pdf(path: Path) -> tuple[str, dict[str, int | bool]]:
    chunks: list[str] = []
    page_count = 0
    table_page_count = 0
    scanned_page_count = 0

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        for idx, page in enumerate(pdf.pages, start=1):
            page_context = f"{path.name}:page:{idx}"
            try:
                table_blocks: list[str] = []
                table_bboxes: list[tuple] = []

                try:
                    tables = page.find_tables() or []
                except Exception as exc:
                    logger.warning("Table detection failed (%s): %s", page_context, exc)
                    tables = []

                for table in tables:
                    try:
                        table_bboxes.append(table.bbox)
                        table_data = table.extract()
                        table_md = _table_to_markdown(table_data)
                        if table_md:
                            table_blocks.append(table_md)
                    except Exception:
                        continue

                page_text = _extract_text_outside_tables(page, table_bboxes, context=page_context)
                if not table_blocks:
                    try:
                        fallback_tables = page.extract_tables() or []
                    except Exception as exc:
                        logger.warning("Fallback table extraction failed (%s): %s", page_context, exc)
                        fallback_tables = []
                    table_blocks = [_table_to_markdown(table) for table in fallback_tables]
                table_blocks = [block for block in table_blocks if block]

                has_table = bool(table_blocks)
                if has_table:
                    table_page_count += 1

                compact_page_text = re.sub(r"\s+", "", page_text)
                is_scanned_like = len(compact_page_text) < 50 and not has_table
                if is_scanned_like:
                    scanned_page_count += 1

                if not page_text.strip() and not has_table:
                    continue

                section_parts = [f"[Page {idx}]"]
                if page_text.strip():
                    section_parts.append(page_text.strip())

                if table_blocks:
                    labeled_tables = [
                        f"[Table {table_idx} on Page {idx}]\n{table_md}"
                        for table_idx, table_md in enumerate(table_blocks, start=1)
                    ]
                    section_parts.append("\n\n".join(labeled_tables))

                chunks.append("\n\n".join(part for part in section_parts if part.strip()))
            except Exception as exc:
                logger.warning("Skipping unreadable PDF page (%s): %s", page_context, exc)
                scanned_page_count += 1
                continue

    return "\n\n".join(chunks), {
        "page_count": page_count,
        "table_page_count": table_page_count,
        "scanned_page_count": scanned_page_count,
        "has_tables": table_page_count > 0,
    }


def _docx_table_to_markdown(table: docx.table.Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells])
    return _table_to_markdown(rows)


def _iter_docx_blocks(doc: docx.document.Document):
    for child in doc.element.body.iterchildren():
        tag = child.tag.lower()
        if tag.endswith("}p"):
            yield "paragraph", child
        elif tag.endswith("}tbl"):
            yield "table", child


def _load_docx(path: Path) -> str:
    doc = docx.Document(path)
    parts: list[str] = []
    table_index = 0

    tables_by_xml = {tbl._tbl: tbl for tbl in doc.tables}
    paras_by_xml = {para._p: para for para in doc.paragraphs}

    for block_type, xml_block in _iter_docx_blocks(doc):
        if block_type == "paragraph":
            para = paras_by_xml.get(xml_block)
            text = re.sub(r"\s+", " ", (para.text if para else "")).strip()
            if text:
                parts.append(text)
            continue

        table = tables_by_xml.get(xml_block)
        if table is None:
            continue
        table_md = _docx_table_to_markdown(table)
        if table_md:
            table_index += 1
            parts.append(f"[Table {table_index}]\n{table_md}")

    return "\n\n".join(parts).strip()


def _load_doc(path: Path) -> str:
    logger.warning("Skipping legacy .doc file '%s'. Convert to .docx for reliable extraction.", path.name)
    return ""


def _load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_documents(docs_dir: str | None = None) -> list[Document]:
    root = Path(docs_dir or settings.DOCS_DIR)
    if not root.exists() or not root.is_dir():
        raise FileNotFoundError(f"Docs directory not found: {root}")

    supported = {".pdf", ".txt", ".md", ".docx", ".doc"}
    documents: list[Document] = []

    for file_path in sorted(root.rglob("*")):
        if not file_path.is_file() or file_path.suffix.lower() not in supported:
            continue

        suffix = file_path.suffix.lower()
        metadata = _derive_source_metadata(file_path, root)

        try:
            if suffix == ".pdf":
                text, pdf_stats = _load_pdf(file_path)
                metadata.update(pdf_stats)
            elif suffix == ".docx":
                text = _load_docx(file_path)
            elif suffix == ".doc":
                text = _load_doc(file_path)
            else:
                text = _load_text(file_path)
        except Exception as exc:
            logger.warning("Skipping file '%s' due extraction error: %s", file_path, exc)
            continue

        text = text.strip()
        if not text:
            continue

        logger.info(
            "Loaded file %s (%d chars, type=%s, topic=%s, subtopic=%s)",
            file_path.name,
            len(text),
            suffix,
            metadata.get("topic", "General"),
            metadata.get("subtopic", ""),
        )
        documents.append(Document(text=text, source=str(file_path), metadata=metadata))

    return documents